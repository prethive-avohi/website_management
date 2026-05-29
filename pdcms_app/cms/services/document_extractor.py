import os
import frappe


def extract_text(file_url: str) -> str:
    """
    Extract plain text from an uploaded file.
    Supports: PDF, DOCX, DOC, TXT.
    """
    file_path = _resolve_path(file_url)
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    if ext in (".docx", ".doc"):
        return _extract_docx(file_path)
    if ext == ".txt":
        return _extract_txt(file_path)

    frappe.throw(f"Unsupported file type: {ext}. Upload PDF, DOCX, or TXT.")


def _resolve_path(file_url: str) -> str:
    if file_url.startswith("/files/"):
        path = frappe.get_site_path("public", "files", file_url.replace("/files/", ""))
    elif file_url.startswith("/private/files/"):
        path = frappe.get_site_path("private", "files", file_url.replace("/private/files/", ""))
    else:
        path = frappe.get_site_path("public", file_url.lstrip("/"))

    if not os.path.exists(path):
        frappe.throw(f"File not found: {path}")
    return path


def _extract_pdf(path: str) -> str:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        frappe.throw("PyPDF2 not installed. Run: bench pip install PyPDF2")

    reader = PdfReader(path)
    parts = [p.extract_text().strip() for p in reader.pages if p.extract_text()]
    if not parts:
        frappe.throw("No text extracted from PDF. File may be image-based.")
    return "\n\n".join(parts)


def _extract_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        frappe.throw("python-docx not installed. Run: bench pip install python-docx")

    doc = Document(path)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.replace("Heading ", "").strip())
            except ValueError:
                level = 2
            parts.append(f"{'#' * level} {text}")
        else:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)

    if not parts:
        frappe.throw("No text extracted from Word document.")
    return "\n\n".join(parts)


def _extract_txt(path: str) -> str:
    with open(path, encoding="utf-8", errors="replace") as f:
        text = f.read().strip()
    if not text:
        frappe.throw("Text file is empty.")
    return text
