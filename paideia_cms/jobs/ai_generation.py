"""
Content generation pipeline — PDF/DOCX → content_json.

Frappe's responsibility:
  1. Extract text from the private file (only Frappe has auth to read it)
  2. Prepare content (smart truncate)
  3. Dispatch to external agent service with extracted text + schema + prompt

All AI logic, all API keys, all model calls live in the agent service.
Frappe never calls OpenAI / Groq / Claude directly.

Result arrives asynchronously via:
  POST /api/method/paideia_cms.api.v1.agent_callback.receive
"""
import json
import os
import traceback

import frappe
import requests

from paideia_cms.logger import get_logger
log = get_logger("ai_generation")


# ---------------------------------------------------------------------------
# Public entry points — called by frappe.enqueue
# ---------------------------------------------------------------------------

def run_page_generation(page_name):
    _run_generation("CMS Page", page_name)


def run_blog_generation(blog_name):
    _run_generation("CMS Blog", blog_name)


def run_course_generation(course_name):
    _run_generation("CMS Course", course_name)


# ---------------------------------------------------------------------------
# Core pipeline
# ---------------------------------------------------------------------------

def _run_generation(doctype, doc_name):
    doc = frappe.get_doc(doctype, doc_name)

    job_log = frappe.get_doc({
        "doctype": "AI Job Log",
        "reference_doctype": doctype,
        "reference_name": doc_name,
        "job_type": "Generate",
        "status": "Queued",
    })
    job_log.insert(ignore_permissions=True)
    frappe.db.commit()

    doc.db_set("ai_job", job_log.name)
    doc.db_set("generation_status", "Queued")

    log.debug("Starting content generation — %s '%s'", doctype, doc_name)

    try:
        settings     = frappe.get_single("CMS Settings")
        agent_url    = (settings.agent_service_url or "").strip().rstrip("/")
        agent_secret = settings.get_password("agent_service_secret", raise_exception=False) or ""

        if not agent_url:
            frappe.throw(
                "Agent Service URL not configured in CMS Settings. "
                "Set it under CMS Settings → Agent Service URL."
            )

        job_log.mark_running()
        doc.db_set("generation_status", "Extracting")

        log.debug("Extracting text from file: %s", doc.source_file)
        raw_text = _extract_text(doc.source_file)
        if not raw_text or not raw_text.strip():
            frappe.throw("Could not extract any text from the uploaded file.")

        log.debug("Extracted %d chars — preparing content", len(raw_text))
        schema, prompt_text = _get_schema_and_prompt(doc)

        provider = (settings.ai_provider or "groq").lower()
        model    = _get_model(settings)
        log.debug("Using provider=%s model=%s", provider, model)

        callback_url = (
            frappe.utils.get_url()
            + "/api/method/paideia_cms.api.v1.agent_callback.receive"
        )

        log.debug("Dispatching to agent service at %s", agent_url)
        resp = requests.post(
            f"{agent_url}/api/v1/generate",
            json={
                "extracted_text": _prepare_content(raw_text),
                "json_schema":    schema,
                "prompt":         prompt_text,
                "provider":       provider,
                "model":          model,
                "webhook_url":    callback_url,
                "context": {
                    "doctype":  doctype,
                    "doc_name": doc_name,
                },
            },
            headers={"X-Agent-Secret": agent_secret},
            timeout=15,
        )
        resp.raise_for_status()
        job_id = resp.json().get("job_id", "")
        log.debug("Agent accepted job — id: %s. Waiting for callback...", job_id)
        doc.db_set("generation_status", f"Generating — agent job {job_id}")
        job_log.db_set("status", "Running")
        frappe.db.commit()

    except Exception:
        err = traceback.format_exc()
        log.warning("Content generation failed for %s '%s': %s",
                    doctype, doc_name, err.strip().splitlines()[-1])
        doc.db_set("generation_status", "Failed")
        doc.db_set("ai_error_log", err[:10000])
        job_log.mark_failed(err)
        frappe.db.commit()
        raise


# ---------------------------------------------------------------------------
# Text extraction — stays in Frappe (private file auth)
# Supported: PDF, DOCX, XLSX, TXT
# ---------------------------------------------------------------------------

_EXTRACTORS = {}  # populated below, after each extractor function is defined


def _extract_text(file_url):
    if not file_url:
        frappe.throw("No source file attached. Upload a PDF, DOCX, XLSX, or TXT file before generating.")

    site_path = frappe.utils.get_site_path()

    if file_url.startswith("/private/files/"):
        rel = file_url.lstrip("/")
    elif file_url.startswith("/files/"):
        rel = "public/" + file_url.lstrip("/")
    else:
        rel = "public/files/" + os.path.basename(file_url)

    full_path = os.path.join(site_path, rel)

    if not os.path.exists(full_path):
        frappe.throw(f"Source file not found on disk: {file_url}")

    ext = os.path.splitext(file_url.lower())[1]
    extractor = _EXTRACTORS.get(ext)
    if not extractor:
        frappe.throw(
            f"Unsupported file type '{ext}'. Upload a PDF, DOCX, XLSX, or TXT file."
        )

    raw = extractor(full_path)
    if not raw or not raw.strip():
        frappe.throw(
            "No extractable text found in the file. "
            "If this is a scanned/image-only PDF, it needs OCR — not currently supported."
        )

    return _deduplicate_text(raw)


def _extract_pdf(path):
    """pdfplumber — extracts text and tables per page (much more reliable than PyPDF2)."""
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    parts.append(page_text)
                for table in page.extract_tables():
                    parts.append(_table_to_text(table))
        return "\n\n".join(parts)
    except Exception as e:
        frappe.throw(f"PDF extraction failed: {e}")


def _extract_docx(path):
    """python-docx — paragraphs AND tables (tables are easy to miss)."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            parts.append(_table_to_text(rows))
        return "\n".join(parts)
    except Exception as e:
        frappe.throw(f"DOCX extraction failed: {e}")


def _extract_xlsx(path):
    """openpyxl — every sheet, every row, formulas resolved to their last computed value."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        parts = []
        for sheet in wb.worksheets:
            rows = [
                [str(cell) for cell in row if cell is not None]
                for row in sheet.iter_rows(values_only=True)
            ]
            rows = [r for r in rows if r]
            if rows:
                parts.append(f"# Sheet: {sheet.title}")
                parts.append(_table_to_text(rows))
        return "\n\n".join(parts)
    except Exception as e:
        frappe.throw(f"XLSX extraction failed: {e}")


def _extract_txt(path):
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    except Exception as e:
        frappe.throw(f"TXT extraction failed: {e}")


_EXTRACTORS.update({
    ".pdf":  _extract_pdf,
    ".docx": _extract_docx,
    ".xlsx": _extract_xlsx,
    ".txt":  _extract_txt,
})


def _table_to_text(rows) -> str:
    """Render a list-of-rows table as pipe-separated text — readable by both humans and AI."""
    return "\n".join(" | ".join(cell or "" for cell in row) for row in rows if any(row))


def _deduplicate_text(text: str) -> str:
    result = text
    while "\n\n\n" in result:
        result = result.replace("\n\n\n", "\n\n")
    return result.strip()


def _prepare_content(text: str, max_chars: int = 50_000) -> str:
    """
    Truncate at a clean paragraph/sentence boundary.
    Default 50K chars (~12,500 tokens) — well within GPT-4o's 128K context window.
    HTTP payload of 50KB is fine; concern only arises at hundreds of KB.
    For truly huge documents (100+ pages), a future S3 text-reference approach
    should replace this — but for typical business docs (20–25K chars), this covers it.
    """
    if len(text) <= max_chars:
        return text

    log.debug("Document truncated: %d chars → %d chars", len(text), max_chars)
    chunk  = text[:max_chars]
    cutoff = int(max_chars * 0.70)
    para   = chunk.rfind("\n\n")
    if para >= cutoff:
        return chunk[:para].strip()
    for end in (". ", ".\n", "! ", "? "):
        pos = chunk.rfind(end)
        if pos >= cutoff:
            return chunk[: pos + 1].strip()
    return chunk[: chunk.rfind(" ")].strip()


# ---------------------------------------------------------------------------
# Schema + prompt — reads from CMS Template (no fallback)
# ---------------------------------------------------------------------------

def _get_schema_and_prompt(doc):
    if not doc.template:
        frappe.throw(
            "No template selected. Select an Active CMS Template before generating."
        )
    template = frappe.get_doc("CMS Template", doc.template)

    schema = template.get_schema()
    if not schema:
        frappe.throw(
            f"Template '{doc.template}' has no JSON Schema. "
            "Add a JSON Schema in the CMS Template before generating."
        )

    prompt_text = template.ai_prompt.strip() if template.ai_prompt else ""
    if not prompt_text:
        prompts = template.get_prompts()
        if prompts:
            prompt_text = frappe.get_doc("CMS Prompt", prompts[0]["name"]).prompt_text.strip()

    if not prompt_text:
        frappe.throw(
            f"Template '{doc.template}' has no AI Prompt. "
            "Add an AI Prompt in the CMS Template before generating."
        )

    return schema, prompt_text


def _get_model(settings) -> str:
    provider = (settings.ai_provider or "groq").lower()
    return {
        "groq":        settings.groq_model or "llama-3.3-70b-versatile",
        "openai":      settings.openai_model or "gpt-4o-mini",
        "claude":      settings.claude_model or "claude-sonnet-4-6",
        "ollama":      settings.ollama_model or "mistral",
        "huggingface": settings.hf_model or "",
    }.get(provider, "gpt-4o-mini")
