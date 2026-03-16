import os
import frappe


def extract_text_from_file(file_url):
    """Extract text content from an uploaded PDF or Word document.

    Args:
        file_url: Frappe file URL (e.g. /files/document.pdf)

    Returns:
        Extracted text as a string.
    """
    file_path = get_file_path(file_url)
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_from_docx(file_path)
    else:
        frappe.throw(f"Unsupported file type: {ext}. Please upload a PDF or Word (.docx) file.")


def get_file_path(file_url):
    """Resolve a Frappe file URL to an absolute file path."""
    if file_url.startswith("/files/"):
        site_path = frappe.get_site_path("public", "files", file_url.replace("/files/", ""))
    elif file_url.startswith("/private/files/"):
        site_path = frappe.get_site_path("private", "files", file_url.replace("/private/files/", ""))
    else:
        site_path = frappe.get_site_path("public", file_url.lstrip("/"))

    if not os.path.exists(site_path):
        frappe.throw(f"File not found at: {site_path}")

    return site_path


def extract_from_pdf(file_path):
    """Extract text from a PDF file using PyPDF2."""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        frappe.throw(
            "PyPDF2 is not installed. Run: bench pip install PyPDF2"
        )

    reader = PdfReader(file_path)
    text_parts = []

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text.strip())

    if not text_parts:
        frappe.throw("Could not extract any text from the PDF. The file may be scanned/image-based.")

    return "\n\n".join(text_parts)


def extract_from_docx(file_path):
    """Extract text from a Word (.docx) file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        frappe.throw(
            "python-docx is not installed. Run: bench pip install python-docx"
        )

    doc = Document(file_path)
    text_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading hierarchy
            if para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").strip()
                try:
                    level = int(level)
                except ValueError:
                    level = 2
                text_parts.append(f"{'#' * level} {text}")
            else:
                text_parts.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    if not text_parts:
        frappe.throw("Could not extract any text from the Word document.")

    return "\n\n".join(text_parts)
